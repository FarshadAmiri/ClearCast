"""
تبدیل فایل‌های Markdown به Word (DOCX)
این اسکریپت از کتابخانه python-docx استفاده می‌کند
"""

import os
from pathlib import Path
import subprocess
import sys

def check_pandoc():
    """بررسی نصب بودن Pandoc"""
    try:
        result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True)
        if result.returncode == 0:
            print("✓ Pandoc نصب شده است")
            return True
    except FileNotFoundError:
        pass
    
    print("⚠ Pandoc نصب نیست")
    return False

def install_pandoc_instructions():
    """راهنمای نصب Pandoc"""
    print("\n" + "="*60)
    print("برای تبدیل Markdown به Word، Pandoc نیاز است")
    print("="*60)
    print("\nروش‌های نصب:\n")
    
    print("1️⃣ Windows (با Chocolatey):")
    print("   choco install pandoc")
    
    print("\n2️⃣ Windows (دستی):")
    print("   - دانلود از: https://pandoc.org/installing.html")
    print("   - نصب و افزودن به PATH")
    
    print("\n3️⃣ Linux:")
    print("   sudo apt install pandoc")
    
    print("\n4️⃣ macOS:")
    print("   brew install pandoc")
    
    print("\nپس از نصب، دوباره این اسکریپت را اجرا کنید.")
    print("="*60 + "\n")

def convert_with_pandoc(md_file, output_file):
    """تبدیل با استفاده از Pandoc"""
    try:
        cmd = ['pandoc', str(md_file), '-o', str(output_file)]
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        return True
    except subprocess.CalledProcessError as e:
        print(f"خطا در تبدیل {md_file}: {e.stderr}")
        return False
    except Exception as e:
        print(f"خطا: {e}")
        return False

def try_markdown_conversion(md_file, output_file):
    """تلاش برای تبدیل با روش‌های جایگزین"""
    try:
        # روش 1: استفاده از markdown2
        try:
            import markdown2
            from docx import Document
            from docx.shared import Pt, Inches
            
            print(f"   در حال تبدیل {md_file.name} با markdown2...")
            
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # تبدیل Markdown به HTML
            html = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks'])
            
            # ایجاد سند Word
            doc = Document()
            
            # تنظیم راست‌چین برای فارسی
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            
            # افزودن محتوا (ساده)
            # این یک پیاده‌سازی ساده است - Pandoc بهتر است
            doc.add_paragraph(md_content)
            
            doc.save(str(output_file))
            return True
            
        except ImportError:
            pass
        
        # روش 2: کپی ساده متن
        print(f"   ایجاد فایل Word ساده از {md_file.name}...")
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            with open(md_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            doc = Document()
            
            # تنظیم فونت پیش‌فرض
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Vazir'
            font.size = Pt(12)
            
            for line in lines:
                if line.strip():
                    # شناسایی headings
                    if line.startswith('# '):
                        p = doc.add_heading(line[2:].strip(), level=1)
                    elif line.startswith('## '):
                        p = doc.add_heading(line[3:].strip(), level=2)
                    elif line.startswith('### '):
                        p = doc.add_heading(line[4:].strip(), level=3)
                    elif line.startswith('#### '):
                        p = doc.add_heading(line[5:].strip(), level=4)
                    elif line.startswith('```'):
                        continue  # Skip code blocks markers
                    else:
                        p = doc.add_paragraph(line.strip())
                        # راست‌چین کردن برای فارسی
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            doc.save(str(output_file))
            return True
            
        except ImportError:
            print("   ⚠ کتابخانه python-docx نصب نیست")
            print("   برای نصب: pip install python-docx")
            return False
            
    except Exception as e:
        print(f"   خطا: {e}")
        return False

def main():
    """تابع اصلی"""
    print("\n" + "="*60)
    print("🔄 تبدیل مستندات Markdown به Word")
    print("="*60 + "\n")
    
    # مسیر پوشه Instructions
    instructions_dir = Path(__file__).parent
    
    # فایل‌های Markdown
    md_files = [
        'شناسنامه_محصول.md',
        'راهنمای_API.md',
        'راهنمای_نصب_و_استفاده.md'
    ]
    
    # بررسی Pandoc
    has_pandoc = check_pandoc()
    
    if not has_pandoc:
        install_pandoc_instructions()
        
        print("آیا می‌خواهید با روش جایگزین (ساده‌تر) ادامه دهید؟ (y/n): ", end='')
        choice = input().strip().lower()
        
        if choice != 'y':
            print("\nلطفاً ابتدا Pandoc را نصب کنید.")
            return
    
    # تبدیل فایل‌ها
    print("\nشروع تبدیل فایل‌ها...\n")
    success_count = 0
    
    for md_file_name in md_files:
        md_file = instructions_dir / md_file_name
        
        if not md_file.exists():
            print(f"⚠ فایل {md_file_name} پیدا نشد")
            continue
        
        # نام فایل خروجی
        output_file = md_file.with_suffix('.docx')
        
        print(f"📄 {md_file_name} → {output_file.name}")
        
        if has_pandoc:
            # استفاده از Pandoc (بهترین کیفیت)
            if convert_with_pandoc(md_file, output_file):
                print(f"   ✓ تبدیل موفق با Pandoc")
                success_count += 1
            else:
                print(f"   ✗ خطا در تبدیل")
        else:
            # استفاده از روش جایگزین
            if try_markdown_conversion(md_file, output_file):
                print(f"   ✓ تبدیل موفق (روش ساده)")
                success_count += 1
            else:
                print(f"   ✗ خطا در تبدیل")
        
        print()
    
    # خلاصه نتایج
    print("="*60)
    print(f"✓ {success_count} از {len(md_files)} فایل با موفقیت تبدیل شدند")
    print("="*60)
    
    if success_count > 0:
        print(f"\nفایل‌های Word در پوشه زیر ذخیره شدند:")
        print(f"📁 {instructions_dir}")
    
    if not has_pandoc and success_count < len(md_files):
        print("\n💡 نکته: برای کیفیت بهتر، Pandoc را نصب کنید")
    
    print("\n✅ پایان")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠ لغو شد توسط کاربر")
    except Exception as e:
        print(f"\n❌ خطای غیرمنتظره: {e}")
        import traceback
        traceback.print_exc()
