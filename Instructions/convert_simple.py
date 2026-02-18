"""
تبدیل خودکار فایل‌های Markdown به Word
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def setup_rtl_document():
    """ایجاد سند Word با پشتیبانی از فارسی"""
    doc = Document()
    
    # تنظیم فونت پیش‌فرض
    style = doc.styles['Normal']
    font = style.font
    font.name = 'B Nazanin'  # یا Tahoma, Arial
    font.size = Pt(11)
    
    # ایجاد یا تنظیم استایل‌های Heading
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
        except KeyError:
            heading_style = doc.styles.add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
        
        heading_font = heading_style.font
        heading_font.name = 'B Nazanin'
        heading_font.bold = True
        heading_font.size = Pt(18 - i * 2)
        heading_font.color.rgb = RGBColor(0, 0, 139)  # آبی تیره
    
    return doc

def parse_markdown_line(line):
    """تحلیل یک خط Markdown"""
    line = line.strip()
    
    # Heading سطح 1
    if line.startswith('# '):
        return ('heading1', line[2:].strip())
    
    # Heading سطح 2
    elif line.startswith('## '):
        return ('heading2', line[3:].strip())
    
    # Heading سطح 3
    elif line.startswith('### '):
        return ('heading3', line[4:].strip())
    
    # Heading سطح 4
    elif line.startswith('#### '):
        return ('heading4', line[5:].strip())
    
    # Code block
    elif line.startswith('```'):
        return ('code_marker', line)
    
    # List item
    elif line.startswith('- ') or line.startswith('* '):
        return ('list', line[2:].strip())
    
    # Numbered list
    elif re.match(r'^\d+\. ', line):
        return ('numbered_list', re.sub(r'^\d+\. ', '', line))
    
    # Table separator
    elif re.match(r'^\|[-:\s|]+\|$', line):
        return ('table_sep', line)
    
    # Table row
    elif line.startswith('|') and line.endswith('|'):
        return ('table_row', line)
    
    # Horizontal rule
    elif line.startswith('---') or line.startswith('***'):
        return ('hr', '')
    
    # Blockquote
    elif line.startswith('>'):
        return ('quote', line[1:].strip())
    
    # Empty line
    elif not line:
        return ('empty', '')
    
    # Normal paragraph
    else:
        return ('paragraph', line)

def add_to_document(doc, line_type, content):
    """افزودن محتوا به سند Word"""
    
    if line_type == 'heading1':
        p = doc.add_heading(content, level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'heading2':
        p = doc.add_heading(content, level=2)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'heading3':
        p = doc.add_heading(content, level=3)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'heading4':
        p = doc.add_heading(content, level=4)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'paragraph':
        if content:
            p = doc.add_paragraph(content)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'list':
        p = doc.add_paragraph(content, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'numbered_list':
        p = doc.add_paragraph(content, style='List Number')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    elif line_type == 'code_marker':
        pass  # Skip code markers for now
    
    elif line_type == 'table_row':
        # جدول را ساده نگه می‌داریم
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        font = p.runs[0].font if p.runs else None
        if font:
            font.name = 'Courier New'
            font.size = Pt(9)
    
    elif line_type == 'quote':
        p = doc.add_paragraph(content)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.left_indent = Inches(0.5)
        if p.runs:
            p.runs[0].italic = True
    
    elif line_type == 'hr':
        doc.add_paragraph('_' * 60)
    
    elif line_type == 'empty':
        doc.add_paragraph()

def convert_markdown_to_docx(md_file, output_file):
    """تبدیل فایل Markdown به Word"""
    try:
        print(f"   در حال خواندن {md_file.name}...")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        print(f"   ایجاد سند Word...")
        doc = setup_rtl_document()
        
        in_code_block = False
        code_lines = []
        
        for line in lines:
            line_type, content = parse_markdown_line(line)
            
            # مدیریت code blocks
            if line_type == 'code_marker':
                if in_code_block:
                    # پایان code block
                    if code_lines:
                        code_text = '\n'.join(code_lines)
                        p = doc.add_paragraph(code_text)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        if p.runs:
                            p.runs[0].font.name = 'Courier New'
                            p.runs[0].font.size = Pt(9)
                        code_lines = []
                    in_code_block = False
                else:
                    # شروع code block
                    in_code_block = True
                continue
            
            if in_code_block:
                code_lines.append(line.rstrip())
                continue
            
            # افزودن به سند
            add_to_document(doc, line_type, content)
        
        print(f"   ذخیره فایل Word...")
        doc.save(str(output_file))
        
        return True
        
    except Exception as e:
        print(f"   ❌ خطا: {e}")
        return False

def main():
    """تابع اصلی"""
    print("\n" + "="*60)
    print("📝 تبدیل خودکار مستندات به Word")
    print("="*60 + "\n")
    
    instructions_dir = Path(__file__).parent
    
    md_files = [
        'شناسنامه_محصول.md',
        'راهنمای_API.md',
        'راهنمای_نصب_و_استفاده.md'
    ]
    
    success_count = 0
    
    for md_file_name in md_files:
        md_file = instructions_dir / md_file_name
        
        if not md_file.exists():
            print(f"⚠ فایل {md_file_name} پیدا نشد\n")
            continue
        
        output_file = md_file.with_suffix('.docx')
        
        print(f"📄 {md_file_name}")
        print(f"   → {output_file.name}")
        
        if convert_markdown_to_docx(md_file, output_file):
            print(f"   ✅ تبدیل موفق\n")
            success_count += 1
        else:
            print(f"   ❌ تبدیل ناموفق\n")
    
    print("="*60)
    print(f"✅ {success_count} از {len(md_files)} فایل تبدیل شدند")
    print("="*60)
    
    if success_count > 0:
        print(f"\n📁 فایل‌های Word در پوشه Instructions ذخیره شدند")
        print(f"   {instructions_dir.absolute()}")
    
    print("\n💡 نکته: برای کیفیت بهتر، از Pandoc استفاده کنید:")
    print("   pandoc فایل.md -o فایل.docx")
    print("\n✅ پایان\n")

if __name__ == "__main__":
    main()
